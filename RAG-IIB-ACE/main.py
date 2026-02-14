import streamlit as st
import os
import time
import re
from google import genai
from google.genai import types
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from io import BytesIO
import instructions  # Ensure instructions.py is in the same folder

# --- 1. SETUP & CONFIGURATION ---
st.set_page_config(page_title="IIB/ACE Code Explainer", page_icon="🤖", layout="wide")

# --- 2. SIDEBAR CONFIGURATION (API KEY INPUT) ---
with st.sidebar:
    st.header("🔑 Configuration")
    
    # Secure Input for API Key
    user_api_key = st.text_input(
        "Enter Google Gemini API Key",
        type="password",
        help="Get your key from aistudio.google.com"
    )
    
    if not user_api_key:
        st.warning("⚠️ Please enter your API Key to proceed.")

# --- 3. CACHED CLIENT INITIALIZATION ---
@st.cache_resource
def get_gemini_client(api_key):
    """
    Creates the Gemini Client only once per API Key.
    """
    return genai.Client(api_key=api_key)

# Initialize Client only if Key is provided
client = None
if user_api_key:
    try:
        client = get_gemini_client(user_api_key)
    except Exception as e:
        st.error(f"Invalid API Key: {e}")

# Initialize Session State (Memory)
if "messages" not in st.session_state:
    st.session_state.messages = []

# Initialize Chat Session (Only if client exists)
if client and "chat_session" not in st.session_state:
    st.session_state.chat_session = client.chats.create(
        model="gemini-2.5-flash",
        config=types.GenerateContentConfig(
            system_instruction=instructions.SYSTEM_PROMPT,
            temperature=0.3,
        )
    )

# --- 4. ADVANCED WORD DOCUMENT GENERATION ---
def add_markdown_paragraph(doc, text, style='Normal'):
    """Helper: Parses **bold** text and adds it to the Word doc."""
    p = doc.add_paragraph(style=style)
    parts = text.split('**')
    for i, part in enumerate(parts):
        run = p.add_run(part)
        if i % 2 == 1:
            run.bold = True
    return p

def create_word_docx(history):
    """Generates a professionally formatted Word document from chat history."""
    doc = Document()
    
    # -- Document Title --
    heading = doc.add_heading('IIB Code Analysis Report', 0)
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph(f"Generated on: {time.strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph("_" * 50)

    for msg in history:
        role = msg["role"]
        content = msg["content"]
        
        # --- USER SECTION ---
        if role == "user":
            h = doc.add_heading('User Query / Code:', level=2)
            h.style.font.color.rgb = RGBColor(0, 51, 102) # Dark Blue
            
            if "CREATE" in content or "SET" in content or len(content.split('\n')) > 3:
                p = doc.add_paragraph()
                run = p.add_run(content)
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
                p.paragraph_format.left_indent = Inches(0.2)
                p.paragraph_format.space_after = Pt(12)
            else:
                p = doc.add_paragraph(content)
                p.italic = True
                p.paragraph_format.space_after = Pt(12)

        # --- ASSISTANT SECTION ---
        elif role == "assistant":
            h = doc.add_heading('Analysis & Explanation:', level=2)
            h.style.font.color.rgb = RGBColor(34, 139, 34) # Forest Green
            
            lines = content.split('\n')
            in_code_block = False
            
            for line in lines:
                line = line.strip()
                if line.startswith("```"):
                    in_code_block = not in_code_block
                    continue
                
                if in_code_block:
                    p = doc.add_paragraph()
                    run = p.add_run(line)
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(0, 100, 0)
                    p.paragraph_format.left_indent = Inches(0.5)
                    p.paragraph_format.space_after = Pt(0)
                    continue
                
                if line.startswith("### "):
                    doc.add_heading(line.replace("### ", ""), level=3)
                elif line.startswith("## "):
                    doc.add_heading(line.replace("## ", ""), level=2)
                elif line.startswith("- ") or line.startswith("* "):
                    add_markdown_paragraph(doc, line[2:], style='List Bullet')
                elif re.match(r'^\d+\.', line):
                    clean_line = line.split('.', 1)[1].strip()
                    add_markdown_paragraph(doc, clean_line, style='List Number')
                elif line:
                    add_markdown_paragraph(doc, line, style='Normal')

            doc.add_paragraph("_" * 50)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# --- 5. STREAMLIT UI LAYOUT ---
st.title("🤖 Aetna ACE Code Explainer")

if not user_api_key:
    st.info("👈 Please enter your **Google Gemini API Key** in the sidebar to start.")
else:
    st.markdown("Paste your **ESQL, Java, or Message Flow XML** below. I will explain the logic step-by-step.")

# Sidebar for Actions
with st.sidebar:
    st.header("📄 Actions")
    
    if st.session_state.messages:
        docx_file = create_word_docx(st.session_state.messages)
        st.download_button(
            label="📥 Download Report (.docx)",
            data=docx_file,
            file_name=f"IIB_Analysis_{int(time.time())}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    
    if st.button("🗑️ Clear Chat"):
        st.session_state.messages = []
        if client:
            st.session_state.chat_session = client.chats.create(
                model="gemini-2.0-flash",
                config=types.GenerateContentConfig(
                    system_instruction=instructions.SYSTEM_PROMPT,
                    temperature=0.3,
                )
            )
        st.rerun()

# --- 6. CHAT INTERFACE ---
for message in st.session_state.messages:
    role = message["role"]
    content = message["content"]
    with st.chat_message(role):
        st.markdown(content)

# Disable chat input if no API key
if user_api_key:
    if prompt := st.chat_input("Paste your IIB Code or ask a question..."):
        st.session_state.messages.append({"role": "user", "content": prompt})
        with st.chat_message("user"):
            st.markdown(prompt)

        with st.chat_message("assistant"):
            with st.spinner("Analyzing logic..."):
                try:
                    response = st.session_state.chat_session.send_message(prompt)
                    st.markdown(response.text)
                    st.session_state.messages.append({"role": "assistant", "content": response.text})
                except Exception as e:
                    st.error(f"Error communicating with Gemini: {e}")
                    st.warning("🔄 Try checking your API Key or clicking 'Clear Chat'.")